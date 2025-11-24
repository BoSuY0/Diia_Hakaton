from pathlib import Path

from docx import Document

from backend.domain.documents.docx_filler import fill_docx_template


def test_fill_docx_replaces_placeholders_and_aliases(tmp_path):
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Name: {{lessor.name}}")
    doc.add_paragraph("Passport: [[lessor.id_doc]]")
    doc.add_paragraph("IBAN: {{lessor.iban}}")
    doc.save(template)

    output = tmp_path / "out.docx"
    fill_docx_template(
        template_path=template,
        field_values={
            "lessor.name": "John",
            "lessor.passport": "AA123456",
            "lessor.iban": "",
        },
        output_path=output,
        keep_placeholders=False,
    )

    result = Document(output)
    text = "\n".join(p.text for p in result.paragraphs)
    assert "John" in text
    # alias id_doc -> passport should be filled
    assert "AA123456" in text
    # empty IBAN placeholder should be cleaned up (no braces)
    assert "{{lessor.iban}}" not in text


def test_fill_docx_keeps_placeholders_in_preview(tmp_path):
    template = tmp_path / "template_preview.docx"
    doc = Document()
    doc.add_paragraph("Empty: {{lessor.rnokpp}}")
    doc.save(template)

    output = tmp_path / "out_preview.docx"
    fill_docx_template(
        template_path=template,
        field_values={},
        output_path=output,
        keep_placeholders=True,
    )

    result = Document(output)
    text = "\n".join(p.text for p in result.paragraphs)
    # Placeholder remains because keep_placeholders=True
    assert "{{lessor.rnokpp}}" in text or "[[lessor.rnokpp]]" in text
