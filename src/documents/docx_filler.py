from __future__ import annotations

from pathlib import Path
from typing import Dict
import re

from docx import Document  # type: ignore
import asyncio
from src.common.async_utils import run_sync


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def fill_docx_template(
    template_path: Path,
    field_values: Dict[str, str],
    output_path: Path,
    *,
    keep_placeholders: bool = False,
) -> Path:
    return _fill_docx_template_sync(template_path, field_values, output_path, keep_placeholders=keep_placeholders)


def _fill_docx_template_sync(
    template_path: Path,
    field_values: Dict[str, str],
    output_path: Path,
    *,
    keep_placeholders: bool = False,
    log_context: tuple[str, str] | None = None,
) -> Path:
    def _replace_text(text: str) -> str:
        updated = text
        for field_id, value in values.items():
            patterns = [
                rf"{{{{\s*{re.escape(field_id)}\s*}}}}",
                rf"\[\[\s*{re.escape(field_id)}\s*\]\]",
            ]
            for pattern in patterns:
                updated = re.sub(pattern, lambda _: value, updated)

        if not keep_placeholders:
            updated = re.sub(r"\{\{[^}]+\}\}", "", updated)
            updated = re.sub(r"\[\[[^]]]+\]\]", "", updated)
        return updated

    def _replace_in_paragraph(paragraph) -> None:
        # Preserve formatting by replacing text inside runs instead of resetting paragraph.text
        for run in paragraph.runs:
            run.text = _replace_text(run.text)

    if template_path.exists():
        # Розширюємо словник полів синонімами для деяких party-полів,
        # щоб підтримати існуючі плейсхолдери у DOCX:
        #   <role>.id_doc   -> <role>.passport
        #   <role>.id_code  -> <role>.rnokpp
        values: Dict[str, str] = dict(field_values)
        for key, val in list(values.items()):
            if key.endswith(".id_doc"):
                alias = key.replace(".id_doc", ".passport")
                values.setdefault(alias, val)
            if key.endswith(".passport"):
                alias = key.replace(".passport", ".id_doc")
                values.setdefault(alias, val)
            if key.endswith(".id_code"):
                alias = key.replace(".id_code", ".rnokpp")
                values.setdefault(alias, val)

        doc = Document(str(template_path))
        for paragraph in _iter_paragraphs(doc):
            _replace_in_paragraph(paragraph)
    else:
        # Fallback: simple document with fields
        doc = Document()
        doc.add_heading("Автоматично згенерований документ", level=1)
        for field_id, value in field_values.items():
            doc.add_paragraph(f"{field_id}: {value}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


async def fill_docx_template_async(
    template_path: Path,
    field_values: Dict[str, str],
    output_path: Path,
    *,
    keep_placeholders: bool = False,
) -> Path:
    # python-docx is sync; run in threadpool
    return await run_sync(
        _fill_docx_template_sync,
        template_path,
        field_values,
        output_path,
        keep_placeholders=keep_placeholders,
    )
