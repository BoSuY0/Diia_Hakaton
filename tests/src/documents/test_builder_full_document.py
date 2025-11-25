"""Full document build integration tests."""
import pytest
from docx import Document

from backend.domain.documents.builder import build_contract
from backend.infra.persistence.store import get_or_create_session, save_session
from backend.domain.sessions.models import FieldState


@pytest.mark.asyncio
async def test_build_contract_with_all_fields(mock_settings, mock_categories_data):
    """Integration test: build contract with all fields and verify output."""
    session_id = "full_contract_doc"
    s = get_or_create_session(session_id)
    s.category_id = mock_categories_data
    s.template_id = "t1"
    s.role = "lessor"
    s.person_type = "individual"

    # Required contract field
    s.contract_fields["cf1"] = FieldState(status="ok")
    s.all_data["cf1"] = {"current": "Contract Value"}

    # Party fields for both roles
    s.party_types["lessor"] = "individual"
    s.party_types["lessee"] = "individual"
    s.party_fields["lessor"] = {"name": FieldState(status="ok")}
    s.party_fields["lessee"] = {"name": FieldState(status="ok")}
    s.all_data["lessor.name"] = {"current": "Lessor Name"}
    s.all_data["lessee.name"] = {"current": "Lessee Name"}

    save_session(s)

    # Prepare a real DOCX template with placeholders
    template_dir = mock_settings.default_documents_root / mock_categories_data
    template_dir.mkdir(parents=True, exist_ok=True)
    template_path = template_dir / "f1.docx"
    doc = Document()
    doc.add_paragraph("CF: {{cf1}}")
    doc.add_paragraph("Lessor: {{lessor.name}}")
    doc.add_paragraph("Lessee: {{lessee.name}}")
    doc.save(str(template_path))

    result = await build_contract(session_id, "t1")
    output_path = result["file_path"]

    built = Document(output_path)
    text = "\n".join(p.text for p in built.paragraphs)

    assert "Contract Value" in text
    assert "Lessor Name" in text
    assert "Lessee Name" in text
