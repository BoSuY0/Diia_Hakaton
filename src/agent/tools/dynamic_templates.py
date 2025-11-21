import json
import uuid
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt

from src.common.config import settings
from src.common.logging import get_logger

logger = get_logger(__name__)

def create_dynamic_contract_template(
    title: str,
    contract_text: str,
    fields_metadata: List[Dict[str, Any]],
    roles: Dict[str, str]
) -> str:
    """
    Creates a temporary dynamic contract template based on LLM generation.
    
    Args:
        title: The title of the contract (e.g. "Договір дарування").
        contract_text: The full text of the contract containing placeholders like {{field_name}}.
                       The text can contain markdown-like formatting.
        fields_metadata: List of field definitions. Each dict must have:
                         - field: str (key, e.g. "price")
                         - label: str (human readable label)
                         - required: bool
                         - type: str (text, date, number, etc)
                         - party: str (optional, e.g. "party_a" or "party_b" if specific to a side)
        roles: Dict defining roles, e.g. {"party_a": "Дарувальник", "party_b": "Обдаровуваний"}
        
    Returns:
        str: The new template_id (e.g. "dynamic_12345...")
    """
    
    # Generate a unique ID
    template_id = f"dynamic_{uuid.uuid4().hex[:8]}"
    
    # 1. Create DOCX template
    doc = Document()
    
    # Add Title
    heading = doc.add_heading(title, 0)
    heading.alignment = 1  # Center
    
    # Add Body Text
    # We split by newlines to preserve paragraphs
    paragraphs = contract_text.split('\n')
    for p_text in paragraphs:
        p_text = p_text.strip()
        if not p_text:
            continue
        
        # Simple heuristic for headings within text (e.g. "1. Предмет договору")
        if p_text[0].isdigit() and "." in p_text[:5] and len(p_text) < 100:
             doc.add_heading(p_text, level=2)
        else:
            p = doc.add_paragraph(p_text)
            p.style.font.size = Pt(11)

    # Save DOCX
    docx_path = settings.assets_dir / "documents" / "templates" / "dynamic" / f"{template_id}.docx"
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    logger.info(f"Created dynamic DOCX at {docx_path}")

    # 2. Create Metadata JSON
    # We need to construct a structure compatible with our Category/Template system
    
    # Separate fields into "common" (contract_fields) and "party specific"
    contract_fields = []
    party_a_fields = []
    party_b_fields = []
    
    for f in fields_metadata:
        party = f.get("party")
        field_def = {
            "field": f["field"],
            "label": f["label"],
            "required": f.get("required", True),
            "type": f.get("type", "text")
        }
        
        if party == "party_a":
            party_a_fields.append(field_def)
        elif party == "party_b":
            party_b_fields.append(field_def)
        else:
            contract_fields.append(field_def)

    meta_data = {
        "category_id": "dynamic",
        "template_id": template_id,
        "label": title,
        "roles": roles,
        "party_modules": {
            "individual": {
                "fields": [
                    {"field": "name", "label": "ПІБ", "required": True},
                    {"field": "address", "label": "Адреса", "required": True},
                    {"field": "tax_id", "label": "РНОКПП", "required": True}
                ]
            },
             "company": {
                "fields": [
                    {"field": "name", "label": "Назва компанії", "required": True},
                    {"field": "edrpou", "label": "ЄДРПОУ", "required": True},
                    {"field": "address", "label": "Юридична адреса", "required": True},
                    {"field": "director", "label": "Директор", "required": True}
                ]
            }
        },
        # Dynamic fields specific to this contract
        "contract_fields": contract_fields,
        # Extra fields injected into parties if needed (though usually we keep parties standard)
        "party_a_extra_fields": party_a_fields,
        "party_b_extra_fields": party_b_fields
    }
    
    json_path = settings.assets_dir / "meta_data" / "dynamic" / f"{template_id}.json"
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(meta_data, ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info(f"Created dynamic Meta JSON at {json_path}")
    
    return template_id
