from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_custom_template():
    doc = Document()
    
    # Title
    p = doc.add_paragraph("{{ contract_title }}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    
    # City and Date
    p = doc.add_paragraph("м. {{ contract_city }}")
    p.add_run("\t\t\t\t\t\t\t\t{{ contract_date }}")
    
    # Parties
    doc.add_paragraph("Сторона 1: {{ party_a.name }}, що діє на підставі власного волевиявлення, надалі іменується «Сторона 1», та")
    doc.add_paragraph("Сторона 2: {{ party_b.name }}, що діє на підставі власного волевиявлення, надалі іменується «Сторона 2»,")
    doc.add_paragraph("уклали цей Договір про наступне:")
    
    # Subject
    doc.add_heading("1. ПРЕДМЕТ ДОГОВОРУ", level=1)
    doc.add_paragraph("{{ contract_subject }}")
    
    # Terms
    doc.add_heading("2. УМОВИ ДОГОВОРУ", level=1)
    doc.add_paragraph("{{ contract_terms }}")
    
    # Signatures
    doc.add_heading("3. РЕКВІЗИТИ ТА ПІДПИСИ СТОРІН", level=1)
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Сторона 1'
    hdr_cells[1].text = 'Сторона 2'
    
    row_cells = table.add_row().cells
    
    # Party A details
    p1 = row_cells[0].paragraphs[0]
    p1.add_run("{{ party_a.name }}\n")
    p1.add_run("Адреса: {{ party_a.address }}\n")
    p1.add_run("ІПН/ЄДРПОУ: {{ party_a.id_code }}\n")
    p1.add_run("IBAN: {{ party_a.iban }}\n")
    p1.add_run("\n_____________________ (Підпис)")
    
    # Party B details
    p2 = row_cells[1].paragraphs[0]
    p2.add_run("{{ party_b.name }}\n")
    p2.add_run("Адреса: {{ party_b.address }}\n")
    p2.add_run("ІПН/ЄДРПОУ: {{ party_b.id_code }}\n")
    p2.add_run("IBAN: {{ party_b.iban }}\n")
    p2.add_run("\n_____________________ (Підпис)")
    
    # Ensure directory exists
    output_dir = "assets/documents_files/default_documents_files"
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = os.path.join(output_dir, "custom_contract.docx")
    doc.save(output_path)
    print(f"Created template at {output_path}")

if __name__ == "__main__":
    create_custom_template()
